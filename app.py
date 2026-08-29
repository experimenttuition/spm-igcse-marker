import json
import io
import time
import streamlit as st
from google import genai
from google.genai import types
import pdfplumber
import docx
from PIL import Image
import pypdfium2 as pdfium

# Initialize Gemini Client using Streamlit Secret
client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])

def process_file(uploaded_file):
    """Extract text or convert images/scanned PDFs into Gemini payload parts."""
    filename = uploaded_file.name.lower()
    payload_parts = []
    
    if filename.endswith(('.png', '.jpg', '.jpeg')):
        img = Image.open(uploaded_file)
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format=img.format if img.format else 'PNG')
        payload_parts.append(("image", types.Part.from_bytes(
            data=img_byte_arr.getvalue(),
            mime_type=uploaded_file.type
        )))
        
    elif filename.endswith('.pdf'):
        # Attempt text extraction first
        text = ""
        with pdfplumber.open(uploaded_file) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
            
        # If PDF has extractable text, use it
        if text.strip():
            payload_parts.append(("text", text))
        else:
            # Fallback for Scanned/Handwritten PDFs: Convert PDF pages to images
            uploaded_file.seek(0)
            pdf_render = pdfium.PdfDocument(uploaded_file.read())
            for page in pdf_render:
                image = page.render(scale=2).to_pil()
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='PNG')
                payload_parts.append(("image", types.Part.from_bytes(
                    data=img_byte_arr.getvalue(),
                    mime_type="image/png"
                )))
                
    elif filename.endswith('.docx'):
        doc = docx.Document(uploaded_file)
        text = "\n".join([p.text for p in doc.paragraphs])
        payload_parts.append(("text", text))
    else:
        payload_parts.append(("text", str(uploaded_file.read(), "utf-8")))
        
    return payload_parts

def generate_text_report(student_name, syllabus, prompt_text, data):
    report = f"WRITING EVALUATION & PARAGRAPH CORRECTION REPORT\n"
    report += f"="*50 + "\n"
    report += f"Student/File : {student_name}\n"
    report += f"Syllabus     : {syllabus}\n"
    report += f"Total Score  : {data.get('overall_score', 0)} / {data.get('max_score', 0)}\n"
    report += f"="*50 + "\n\n"
    
    report += "PARAGRAPH-BY-PARAGRAPH ANALYSIS & DETAILED WORD-LEVEL CORRECTIONS:\n"
    report += "-"*50 + "\n"
    for item in data.get('paragraph_analysis', []):
        report += f"Paragraph {item.get('paragraph_number')}:\n"
        report += f"Original: {item.get('original_text')}\n\n"
        report += "Word-by-Word Explanations & Error Analysis:\n"
        report += f"{item.get('whats_wrong')}\n\n"
        report += f"Corrected Paragraph: {item.get('corrected_text')}\n\n"
        report += "-"*30 + "\n\n"

    report += "CRITERIA BREAKDOWN:\n"
    report += "-"*50 + "\n"
    for item in data.get('breakdown', []):
        report += f"- {item.get('criterion')}: {item.get('score')}/{item.get('max')}\n"
        report += f"  Notes: {item.get('feedback')}\n\n"
        
    report += "KEY STRENGTHS:\n"
    for s in data.get('strengths', []):
        report += f"- {s}\n"
        
    report += "\nAREAS FOR IMPROVEMENT:\n"
    for i in data.get('improvements', []):
        report += f"- {i}\n"
        
    return report

def generate_with_retry(contents_payload, system_instruction, max_retries=3):
    models_to_try = ["gemini-3.6-flash", "gemini-2.5-flash"]
    
    for model_name in models_to_try:
        for attempt in range(max_retries):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=contents_payload,
                    config=types.GenerateContentConfig(
                        system_instruction=system_instruction,
                        response_mime_type="application/json"
                    )
                )
                return response
            except Exception as e:
                if "503" in str(e) and attempt < max_retries - 1:
                    time.sleep(2)
                    continue
                elif attempt == max_retries - 1 and model_name == models_to_try[-1]:
                    raise e

# Interface Setup
st.set_page_config(page_title="Multi-Syllabus Essay Marker", layout="wide")
st.title("📝 Automated Writing Marker & Detailed Paragraph Corrector")

# Expanded Syllabus List
syllabus_list = [
    "MPT4",
    "UASA Form 3",
    "SPM 1119",
    "IGCSE 3138 (Year 7)",
    "IGCSE 3139 (Year 8)",
    "IGCSE 3140 (Year 9)",
    "IGCSE 0816/01",
    "IGCSE 0816/02",
    "IGCSE 0500",
    "IGCSE 0510",
    "IGCSE O-Level 1123"
]

syllabus = st.sidebar.selectbox("Select Syllabus / Marking Scheme", syllabus_list)
task_prompt = st.sidebar.text_area("The Question (Optional)", help="Paste the essay topic or exam question here.")

uploaded_files = st.file_uploader(
    "Upload Student Essay Pages (.png, .jpg, .jpeg, .pdf, .docx, .txt)", 
    type=["png", "jpg", "jpeg", "pdf", "docx", "txt"],
    accept_multiple_files=True
)

if uploaded_files and st.button("Mark & Correct Essay"):
    contents_payload = []
    
    if task_prompt:
        contents_payload.append(f"THE QUESTION / ASSIGNMENT PROMPT:\n{task_prompt}\n\n")

    for idx, file in enumerate(uploaded_files):
        extracted_parts = process_file(file)
        for part_type, content in extracted_parts:
            if part_type == "image":
                contents_payload.append(f"STUDENT ESSAY IMAGE PAGE {idx + 1}:")
                contents_payload.append(content)
            else:
                contents_payload.append(f"STUDENT ESSAY TEXT PAGE {idx + 1}:\n{content}")

    system_instruction = f"""
    You are an official examiner for {syllabus}. 
    Evaluate the provided student essay strictly according to official assessment rubrics and criteria for {syllabus}.
    The essay may span across MULTIPLE uploaded images/pages or scanned PDFs. Read all pages in order as ONE single continuous essay.

    CRITICAL INSTRUCTIONS FOR TRANSCRIPTION AND EXPLANATION:
    1. IGNORE any crossed-out, struck-through, or erased words in the student's text. Treat them as if they were never written.
    2. Provide a granular, highly clear breakdown for EVERY SINGLE WORD or PHRASE you change, delete, or add.
    3. For the `whats_wrong` field, structure the explanation using clear bullet points. For EVERY changed word/phrase, specify:
       - **Original Wording:** [exact word/phrase]
       - **Why Unacceptable:** [explain clearly why it is grammatically incorrect, awkward, informal, misplaced, or unidiomatic for {syllabus}]
       - **Replacement/Correction:** [exact word/phrase used in the corrected version]

    Return JSON ONLY matching this structure:
    {{
      "overall_score": 0,
      "max_score": 0,
      "paragraph_analysis": [
        {{
          "paragraph_number": 1,
          "original_text": "Exact text from paragraph 1 (excluding crossed-out words)",
          "whats_wrong": "Bulleted list providing clear explanations for every changed word and why the original phrasing is unacceptable.",
          "corrected_text": "Polished and corrected version of paragraph 1"
        }}
      ],
      "breakdown": [
        {{"criterion": "Criterion Name", "score": 0, "max": 0, "feedback": "Detailed notes referencing text."}}
      ],
      "strengths": ["Strength 1"],
      "improvements": ["Improvement 1"]
    }}
    """
    
    with st.spinner(f"Analyzing essay using {syllabus} rubric standards..."):
        try:
            response = generate_with_retry(contents_payload, system_instruction)
            data = json.loads(response.text)
            
            st.markdown("---")
            st.header(f"Total Score: {data['overall_score']} / {data['max_score']}")
            
            st.subheader("✍️ Paragraph-by-Paragraph Corrections & Word Explanations")
            for item in data.get('paragraph_analysis', []):
                with st.expander(f"📌 Paragraph {item['paragraph_number']} Analysis & Corrections", expanded=True):
                    st.markdown(f"**Original Text:**\n> {item['original_text']}")
                    st.markdown(f"**🔍 Detailed Word-by-Word Explanations & Error Analysis:**\n{item['whats_wrong']}")
                    st.markdown(f"**✅ Corrected Paragraph:**\n{item['corrected_text']}")

            st.subheader("📊 Criteria Breakdown")
            for item in data.get('breakdown', []):
                st.markdown(f"**{item['criterion']} ({item['score']}/{item['max']})**")
                st.write(item['feedback'])
                
            st.subheader("✅ Key Strengths")
            for s in data.get('strengths', []):
                st.write(f"- {s}")
                
            st.subheader("💡 Areas for Improvement")
            for i in data.get('improvements', []):
                st.write(f"- {i}")
                
            report_name = uploaded_files[0].name.split('.')[0]
            report_str = generate_text_report(report_name, syllabus, task_prompt, data)
            st.download_button(
                label="📥 Download Full Marking & Correction Sheet (.txt)",
                data=report_str,
                file_name=f"Marked_{report_name}.txt",
                mime="text/plain"
            )

        except Exception as e:
            st.error(f"Error evaluating essay: {e}")
